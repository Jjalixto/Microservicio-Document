from models.document_request import DocumentRequest
from strategies.one_owner import OneOwner
from strategies.two_owners import TwoOwners
from strategies.three_owners import ThreeOwners

import xlwings as xw
import openpyxl
from datetime import datetime
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class DocumentoService:
    
    """
    Servicio central para manejar la generación de documentos.
    """
    def servicioCentral(self, request: DocumentRequest):
        
        # Validación del número de propietarios
        if request.propietario not in [1, 2, 3]:
            raise ValueError("El número de propietarios debe ser 1, 2 o 3.")
        # Selección de la estrategia adecuada según el número de propietarios
        
        if request.propietario == 1:
            strategy = OneOwner()
        elif request.propietario == 2:
            strategy = TwoOwners()  
        elif request.propietario == 3:
            strategy = ThreeOwners()
        
        # Procesar la solicitud con la estrategia seleccionada
        return strategy.process_request(request)

<<<<<<< HEAD
=======
            # Fecha
            'day': request.day or '',
            'month': request.month or '',
            'year': request.year or '',

            # Datos del cliente 1
            'name_1': request.name_1 or '',
            'dni_1': request.dni_1 or '',
            'ocupation_1': request.ocupation_1 or '',
            'marital_status_1': request.marital_status_1 or '',
            'address_1': request.address_1 or '',
            'mail_1': request.mail_1 or '',
            'phone_1': request.phone_1 or '',

            # Datos del cliente 2
            'name_2': request.name_2 or '',
            'dni_2': request.dni_2 or '',
            'ocupation_2': request.ocupation_2 or '',
            'marital_status_2': request.marital_status_2 or '',
            'address_2': request.address_2 or '',
            'mail_2': request.mail_2 or '',
            'phone_2': request.phone_2 or '',

            # Datos del lote
            'number_batch': request.number_batch or '',
            'approximate_area': request.approximate_area or '',
        }
        
        document.render(valores)
        
        DocumentoService.reemplazar_marcadores(document, valores)
        DocumentoService.eliminar_parrafos_innecesarios(document)
        DocumentoService.eliminar_desde_marcador(document, '${eliminar}')
        
        ruta_word = './documento_contado_final.docx'
        ruta_excel='lib/calculadora.xlsx'
        
        document.save(ruta_word)
        
        DocumentoService.actualizar_excel(ruta_excel, request)
        
        # Leer datos del archivo Excel
        tabla_datos = DocumentoService.leer_datos_excel(ruta_excel)
        DocumentoService.actualizar_documento_word(ruta_word, tabla_datos)
        
        print("documento word actualizado y guardado")
        
        
    @staticmethod
    def actualizar_excel(ruta_excel, request: DocumentRequest):
        fecha = datetime.strptime(request.fecha_primera_cuota, '%d/%m/%Y')
        app = xw.App(visible=False)
        wb = xw.Book(ruta_excel)
        try:
            hoja1 = wb.sheets['Calculadora']
            hoja1.range('C1').value = request.monto_venta  # Inserta un valor en la celda C1
            hoja1.range('C2').value = request.cuota_inicial   # Inserta un valor en la celda C2
            hoja1.range('C4').value = request.cantidad_anios      # Inserta un valor en la celda C4
            hoja1.range('C5').value = fecha  # Inserta una fecha en la celda C5
            wb.save()
        finally:
            wb.close()
            app.quit()
            
    @staticmethod
    def leer_datos_excel(ruta_archivo):
        """
        Lee los datos calculados del archivo Excel desde la hoja 'Calculadora'.
        """
        workbook = openpyxl.load_workbook(ruta_archivo, data_only=True)
        hoja1 = workbook['Calculadora']

        tabla_datos_hoja1 = []
        fila = 12  # Comenzamos desde la fila 12
        while True:
            celda_b = hoja1[f'B{fila}'].value
            if celda_b is None or celda_b == 0 or celda_b == '':
                break

            fila_datos = [hoja1[f'{col}{fila}'].value for col in ['B', 'C', 'E', 'F', 'G', 'H', 'I', 'J']]
            tabla_datos_hoja1.append(fila_datos)
            fila += 1
        
        return tabla_datos_hoja1
    
    def actualizar_documento_word(ruta_archivo_word, tabla_datos):
        """
        Actualiza el documento Word reemplazando el marcador '${cronograma}' con una tabla.
        """
        doc = Document(ruta_archivo_word)
        DocumentoService.agregar_tabla_word(doc, '${cronograma}', tabla_datos)
        doc.save(ruta_archivo_word)
    
    def agregar_tabla_word(doc, marcador, tabla_datos_hoja1):
        
    # Buscar y reemplazar `${cronograma}` con la tabla en su lugar
        for paragraph in doc.paragraphs:
            if '${cronograma}' in paragraph.text:
                paragraph.text = paragraph.text.replace('${cronograma}', '')  # Quitar la cadena de marcador

                # Insertar la tabla justo después del párrafo
                if tabla_datos_hoja1:
                    # Crear la tabla
                    tabla = doc.add_table(rows=1, cols=len(tabla_datos_hoja1[0]))
                    hdr_cells = tabla.rows[0].cells
                    encabezados = ['Nro Cuota', 'Fecha de Vencimiento', 'Saldo Capital', 'Cuota Capital', 'Cuota Interés', 'Cuota Admin.', 'Cuota ITF', 'Cuota Total']

                    # Insertar encabezados con formato
                    for i, header in enumerate(encabezados):
                        hdr_cells[i].text = header
                        hdr_cells[i]._element.get_or_add_tcPr().append(
                            parse_xml(r'<w:shd {} w:fill="CCFFCC"/>'.format(nsdecls('w')))
                        )
                        for p in hdr_cells[i].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(8)

                    # Insertar filas de datos
                    for fila_datos in tabla_datos_hoja1:
                        row_cells = tabla.add_row().cells
                        for i, valor in enumerate(fila_datos):
                            if i == 1 and isinstance(valor, datetime):  # Formato de fecha
                                row_cells[i].text = valor.strftime('%Y-%m-%d')
                            elif i >= 2:  # Formato de números con dos decimales
                                row_cells[i].text = f"{valor:.2f}" if isinstance(valor, (int, float)) else ''
                            else:
                                row_cells[i].text = str(valor) if valor is not None else ''
                            for p in row_cells[i].paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(8)

                    # Aplicar bordes a la tabla
                    tbl_xml = tabla._tbl
                    tbl_borders = parse_xml(
                        r'<w:tblBorders {}>'
                        r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        r'</w:tblBorders>'.format(nsdecls('w'))
                    )
                    tbl_xml.tblPr.append(tbl_borders)

                    # Mover la tabla al lugar del marcador usando `addnext`
                    paragraph._p.addnext(tabla._element)
    
    @staticmethod
    def financed_type(request: DocumentRequest, document: Document):
        
        valores = {

        #financed according to the table
        # 'texto_1':'Anexo Nº 5: Hoja Resumen',
        # 'texto_2':'El Comprador declara conocer que las indicadas son cuentas recaudadoras razón por la que ante el incumplimiento de pago en la fecha correspondiente incurrirá en mora automática sin necesidad de intimación previa; en consecuencia, se devengará un interés compensatorio diario de US$ 1.00 (Un y 00/100 dólares americanos), y un interés moratorio diario igual, ambos respecto del importe de la cuota adeudada, los cuales se cobrarán conjuntamente con la cuota pendiente de pago. El Comprador reconoce que los pagos deben efectuarse, obligatoriamente, a través de dicha cuenta recaudadora, considerándose esta como una obligación contractual <br> Sin perjuicio de ello, El Comprador declara conocer que, supletoriamente al sistema de recaudación mencionado, podrá realizar el pago de las cuotas mediante el acceso a un enlace de pago generado por la Vendedora y/o sistema de recaudación propuesto por la Vendedora; el mismo que también generará una mora automática compuesta por un interés compensatorio diario y un interés moratorio diario del mismo valor señalado en el párrafo anterior, siempre que incumple con el pago de la cuota en la fecha correspondiente. Las partes declaran que esta forma de pago también se considerara una obligación contractual y generará los efectos cancelatorios correspondientes. <br> Finalmente, el Comprador deberá informar y enviar a la Vendedora, los sustentos de pagos respectivos.',
        # 'texto_3':'Adicionalmente, las partes dejan constancia que, al amparo de lo dispuesto por el artículo 1583 del Código Civil, la Vendedora se reserva la propiedad de el/los lote(s) hasta la cancelación total del Precio de Venta.',
        'texto_4':'La entrega de la posesión de el/los lote(s) se realizará en el mes de diciembre de 202x.',
        'texto_5':'La entrega de la posesión de las áreas y servicios comunes del Condominio se realizará en el mes de diciembre de 202x.',
        # 'texto_6':'La Vendedora podrá reportar a las centrales de riesgo a El Comprador en caso de incumplimiento en el pago de sus cuotas.',
        'texto_7':'(a) dos o más armadas alternas o consecutivas (cuotas) del Precio de Venta adeudado bajo el presente Contrato señaladas en el Cronograma de Pagos indicado en el Numeral 10 del Anexo N.° 5: Hoja Resumen; y/o (b)',
        'texto_8':'Así, en caso el Comprador mantenga algún reclamo que esté siendo materia de controversia no podrá suspender el pago de las cuotas del financiamiento que mantenga pendientes en atención al lote adquirido ni podrá suspender las demás obligaciones que haya contraído, salvo que cuente con una orden judicial o arbitral que así lo determine.',
        'texto_9':'El saldo de US$ --- (--- con 00/100 dólares americanos), que será cancelado',
        'texto_10': 'según el cronograma de pago indicado en el Numeral 10 del Anexo 5: Hoja Resumen', 
        'texto_11': '',
        'texto_12': '',
        
        #day and month
        'day': request.day or '',
        'month': request.month or '',
        'year': request.year or '',

        # Datos del cliente 1
        'name_1': request.name_1 or '',
        'dni_1': request.dni_1 or '',
        'ocupation_1': request.ocupation_1 or '',
        'marital_status_1': request.marital_status_1 or '',
        'address_1': request.address_1 or '',
        'mail_1': request.mail_1 or '',
        'phone_1': request.phone_1 or '',

        # Datos del cliente 2
        'name_2': request.name_2 or '',
        'dni_2': request.dni_2 or '',
        'ocupation_2': request.ocupation_2 or '',
        'marital_status_2': request.marital_status_2 or '',
        'address_2': request.address_2 or '',
        'mail_2': request.mail_2 or '',
        'phone_2': request.phone_2 or '',

        # Datos del lote
        'number_batch': request.number_batch or '',
        'approximate_area': request.approximate_area or '',
        }
        
        document.render(valores)
        
        DocumentoService.insertar_texto_estatico(document)
        DocumentoService.reemplazar_marcadores(document, valores)
        DocumentoService.dejar_el_marcador(document)
        
        ruta_word = './documento_financiado_final.docx'
        ruta_excel='lib/calculadora.xlsx'

        document.save(ruta_word)
        
        DocumentoService.actualizar_excel(ruta_excel, request)
        
        # Leer datos del archivo Excel
        tabla_datos = DocumentoService.leer_datos_excel(ruta_excel)
        DocumentoService.actualizar_documento_word(ruta_word, tabla_datos)
        
        print("documento word actualizado y guardado")
        
    
    @staticmethod
    def fractionated_type(request: DocumentRequest, document: Document):
        valores = {

        #financed according to the table
        # 'texto_2':'El Comprador declara conocer que las indicadas son cuentas recaudadoras razón por la que ante el incumplimiento de pago en la fecha correspondiente incurrirá en mora automática sin necesidad de intimación previa; en consecuencia, se devengará un interés compensatorio diario de US$ 1.00 (Un y 00/100 dólares americanos), y un interés moratorio diario igual, ambos respecto del importe de la cuota adeudada, los cuales se cobrarán conjuntamente con la cuota pendiente de pago. El Comprador reconoce que los pagos deben efectuarse, obligatoriamente, a través de dicha cuenta recaudadora, considerándose esta como una obligación contractual <br> Sin perjuicio de ello, El Comprador declara conocer que, supletoriamente al sistema de recaudación mencionado, podrá realizar el pago de las cuotas mediante el acceso a un enlace de pago generado por la Vendedora y/o sistema de recaudación propuesto por la Vendedora; el mismo que también generará una mora automática compuesta por un interés compensatorio diario y un interés moratorio diario del mismo valor señalado en el párrafo anterior, siempre que incumple con el pago de la cuota en la fecha correspondiente. Las partes declaran que esta forma de pago también se considerara una obligación contractual y generará los efectos cancelatorios correspondientes. <br> Finalmente, el Comprador deberá informar y enviar a la Vendedora, los sustentos de pagos respectivos.',
        # 'texto_3':'Adicionalmente, las partes dejan constancia que, al amparo de lo dispuesto por el artículo 1583 del Código Civil, la Vendedora se reserva la propiedad de el/los lote(s) hasta la cancelación total del Precio de Venta.',
        'texto_4':'La entrega de la posesión de el/los lote(s) se realizará en el mes de diciembre de 202x.',
        'texto_5':'La entrega de la posesión de las áreas y servicios comunes del Condominio se realizará en el mes de diciembre de 202x.',
        # 'texto_6':'La Vendedora podrá reportar a las centrales de riesgo a El Comprador en caso de incumplimiento en el pago de sus cuotas.',
        'texto_7':'(a) dos o más armadas alternas o consecutivas (cuotas) del Precio de Venta adeudado bajo el presente Contrato señaladas en el Cronograma de Pagos indicado en el Numeral 10 del Anexo N.° 5: Hoja Resumen; y/o (b)',
        'texto_8':'Así, en caso el Comprador mantenga algún reclamo que esté siendo materia de controversia no podrá suspender el pago de las cuotas del financiamiento que mantenga pendientes en atención al lote adquirido ni podrá suspender las demás obligaciones que haya contraído, salvo que cuente con una orden judicial o arbitral que así lo determine.',
        'texto_9':'El saldo de US$ --- (--- con 00/100 dólares americanos), que será cancelado',
        'texto_10':'',
        'texto_11': 'según lo siguiente:',
        'texto_12': '(i)	La suma de US$ --- (--- con 00/100 dólares americanos), a más tardar el – de – de 202-. \n (ii)	La suma de US$ --- (--- con 00/100 dólares americanos), a más tardar el – de – de 202-. \n (iii)	(….)',
        
        #day and month
        'day': request.day or '',
        'month': request.month or '',
        'year': request.year or '',

        # Datos del cliente 1
        'name_1': request.name_1 or '',
        'dni_1': request.dni_1 or '',
        'ocupation_1': request.ocupation_1 or '',
        'marital_status_1': request.marital_status_1 or '',
        'address_1': request.address_1 or '',
        'mail_1': request.mail_1 or '',
        'phone_1': request.phone_1 or '',

        # Datos del cliente 2
        'name_2': request.name_2 or '',
        'dni_2': request.dni_2 or '',
        'ocupation_2': request.ocupation_2 or '',
        'marital_status_2': request.marital_status_2 or '',
        'address_2': request.address_2 or '',
        'mail_2': request.mail_2 or '',
        'phone_2': request.phone_2 or '',

        # Datos del lote
        'number_batch': request.number_batch or '',
        'approximate_area': request.approximate_area or '',
        }
        
        document.render(valores)
        
        DocumentoService.insertar_texto_estatico_fractionated(document)
        DocumentoService.reemplazar_marcadores(document, valores)
        DocumentoService.eliminar_parrafos_innecesarios_fractioned(document)
        DocumentoService.eliminar_desde_marcador(document, '${eliminar}')    
        
        document.save('documento_fraccionado_final.docx')
    
    @staticmethod
    def insertar_texto_estatico(document):
        valores= {
        '${texto_1}':'Anexo Nº 5: Hoja Resumen',
        '${texto_2}':'El Comprador declara conocer que las indicadas son cuentas recaudadoras razón por la que ante el incumplimiento de pago en la fecha correspondiente incurrirá en mora automática sin necesidad de intimación previa; en consecuencia, se devengará un interés compensatorio diario de US$ 1.00 (Un y 00/100 dólares americanos), y un interés moratorio diario igual, ambos respecto del importe de la cuota adeudada, los cuales se cobrarán conjuntamente con la cuota pendiente de pago. El Comprador reconoce que los pagos deben efectuarse, obligatoriamente, a través de dicha cuenta recaudadora, considerándose esta como una obligación contractual <br> Sin perjuicio de ello, El Comprador declara conocer que, supletoriamente al sistema de recaudación mencionado, podrá realizar el pago de las cuotas mediante el acceso a un enlace de pago generado por la Vendedora y/o sistema de recaudación propuesto por la Vendedora; el mismo que también generará una mora automática compuesta por un interés compensatorio diario y un interés moratorio diario del mismo valor señalado en el párrafo anterior, siempre que incumple con el pago de la cuota en la fecha correspondiente. Las partes declaran que esta forma de pago también se considerara una obligación contractual y generará los efectos cancelatorios correspondientes. <br> Finalmente, el Comprador deberá informar y enviar a la Vendedora, los sustentos de pagos respectivos.',
        '${texto_3}':'Adicionalmente, las partes dejan constancia que, al amparo de lo dispuesto por el artículo 1583 del Código Civil, la Vendedora se reserva la propiedad de el/los lote(s) hasta la cancelación total del Precio de Venta.',
        '${texto_6}':'La Vendedora podrá reportar a las centrales de riesgo a El Comprador en caso de incumplimiento en el pago de sus cuotas.',
        }
        
        # Iterar sobre todos los párrafos del documento y reemplazar los marcadores
        for paragraph in document.paragraphs:
            for marcador, reemplazo in valores.items():
                if marcador in paragraph.text:  # Verifica si el marcador está en el párrafo
                    # Reemplaza el marcador con el texto proporcionado
                    paragraph.text = paragraph.text.replace(marcador, reemplazo)
    
    @staticmethod
    def insertar_texto_estatico_fractionated(document):
        valores_estaticos = {
        '${texto_2}':'El Comprador declara conocer que las indicadas son cuentas recaudadoras razón por la que ante el incumplimiento de pago en la fecha correspondiente incurrirá en mora automática sin necesidad de intimación previa; en consecuencia, se devengará un interés compensatorio diario de US$ 1.00 (Un y 00/100 dólares americanos), y un interés moratorio diario igual, ambos respecto del importe de la cuota adeudada, los cuales se cobrarán conjuntamente con la cuota pendiente de pago. El Comprador reconoce que los pagos deben efectuarse, obligatoriamente, a través de dicha cuenta recaudadora, considerándose esta como una obligación contractual <br> Sin perjuicio de ello, El Comprador declara conocer que, supletoriamente al sistema de recaudación mencionado, podrá realizar el pago de las cuotas mediante el acceso a un enlace de pago generado por la Vendedora y/o sistema de recaudación propuesto por la Vendedora; el mismo que también generará una mora automática compuesta por un interés compensatorio diario y un interés moratorio diario del mismo valor señalado en el párrafo anterior, siempre que incumple con el pago de la cuota en la fecha correspondiente. Las partes declaran que esta forma de pago también se considerara una obligación contractual y generará los efectos cancelatorios correspondientes. <br> Finalmente, el Comprador deberá informar y enviar a la Vendedora, los sustentos de pagos respectivos.',
        '${texto_3}':'Adicionalmente, las partes dejan constancia que, al amparo de lo dispuesto por el artículo 1583 del Código Civil, la Vendedora se reserva la propiedad de el/los lote(s) hasta la cancelación total del Precio de Venta.',
        '${texto_6}':'La Vendedora podrá reportar a las centrales de riesgo a El Comprador en caso de incumplimiento en el pago de sus cuotas.',
        }
        
        # Iterar sobre todos los párrafos del documento y reemplazar los marcadores
        for paragraph in document.paragraphs:
            for marcador, reemplazo in valores_estaticos.items():
                if marcador in paragraph.text:  # Verifica si el marcador está en el párrafo
                    # Reemplaza el marcador con el texto proporcionado
                    paragraph.text = paragraph.text.replace(marcador, reemplazo)
    
    # Función para reemplazar los marcadores con los valores correspondientes en el documento
    @staticmethod
    def reemplazar_marcadores(document, valores):
        for paragraph in document.paragraphs:
            for key, value in valores.items():
                if key in paragraph.text:
                    # Reemplazar el marcador con el valor
                    inline = paragraph.runs
                    for run in inline:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
    # Función para eliminar todos los párrafos después de un marcador
    @staticmethod
    def eliminar_desde_marcador(document, marcador):
        eliminar = False
        for paragraph in document.paragraphs:
            if marcador in paragraph.text:
                eliminar = True  # Cuando encontramos el marcador, comenzamos a eliminar
            if eliminar:
                # Elimina el párrafo
                p = paragraph._element
                p.getparent().remove(p)

    @staticmethod
    def dejar_el_marcador(document):
        
        valores = {
            '${eliminar}': ''
        }
        
        # Iterar sobre todos los párrafos del documento y reemplazar los marcadores
        for paragraph in document.paragraphs:
            for marcador, reemplazo in valores.items():
                if marcador in paragraph.text:  # Verifica si el marcador está en el párrafo
                    # Reemplaza el marcador con el texto proporcionado
                    paragraph.text = paragraph.text.replace(marcador, reemplazo)
        
    # Función para eliminar los párrafos que contienen ciertos marcadores
    @staticmethod
    def eliminar_parrafos_innecesarios(document):
        for paragraph in document.paragraphs:
            if '${texto_1}' in paragraph.text or '${texto_2}' in paragraph.text or '${texto_3}' in paragraph.text or '${texto_6}' in paragraph.text:
                p = paragraph._element
                p.getparent().remove(p)

    @staticmethod
    def eliminar_parrafos_innecesarios_fractioned(document):
        for paragraph in document.paragraphs:
            if '${texto_1}' in paragraph.text:
                p = paragraph._element               
                p.getparent().remove(p)
>>>>>>> ca850361e9423ef1761411922acf5c14a9e5d2c7
